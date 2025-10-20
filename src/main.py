import warnings
import os
import argparse
import pathlib
import tempfile
import json
import logging 

# File Handling & Text Processing
import pdfplumber
import docx

# LangChain / OpenAI (updated for LangChain 0.2+)
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import SimpleJsonOutputParser
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI

# Utility
from dotenv import load_dotenv

# Set logging and warnings
logging.getLogger("pdfminer").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", category=UserWarning, module="pdfminer")

# Streamlit setup
try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ImportError:
    STREAMLIT_AVAILABLE = False

# Configuration & Initialization
#load_dotenv(os.path.expanduser("/home/es/Desktop/code/.env")) # Load API key
openai_api_key = st.secrets["OPENAI_API_KEY"] 

MATCHING_PROMPT_TEMPLATE = """
You are an expert recruiter. Your task is to compare a candidate's resume against a job description.
The goal is to determine the suitability and quantify the match.

Return your complete response as a single JSON object that conforms to this schema:
{{
    "score": <int, 0-100, where 100 is a perfect match>,
    "explanation": <list of strings, key bullet points justifying the score, highlighting matches and gaps>
}}

---
Resume:
{resume}

---
Job Description:
{job}

---
Retrieved Context (from similar documents):
{context}
"""

# File Handling and Text Extraction
def extract_and_clean(path: str) -> str:
    """Extracts text content from PDF, DOCX, or plain text files and cleans it."""
    p = pathlib.Path(path)
    suffix = p.suffix.lower()
    text = ""

    try:
        if suffix == ".pdf":
            with pdfplumber.open(path) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif suffix in (".docx", ".doc"):
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            text = p.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"Error extracting text from {path}: {e}")
        return ""

    # Clean the extracted text: remove empty lines and strip whitespace
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())

# LangChain Core Components
def get_embeddings() -> OpenAIEmbeddings:
    """Initializes and returns the OpenAI Embeddings model."""
    print("[Info] Initializing OpenAI embeddings...")
    return OpenAIEmbeddings()

def index_documents(data_dir: str, index_dir: str):
    """Loads documents, embeds them, and saves the FAISS index."""
    resumes_dir = pathlib.Path(data_dir) / "resumes"
    jds_dir = pathlib.Path(data_dir) / "jds"
    docs: list[Document] = []
    valid_extensions = {".pdf", ".txt", ".docx", ".doc"}

    for folder, source in [(resumes_dir, "resume"), (jds_dir, "job")]:
        for f in folder.glob("**/*"):
            if f.suffix.lower() in valid_extensions:
                text = extract_and_clean(str(f)) 
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"source": source, "filename": f.name}))

    if not docs:
        print(f"[Warning] No documents found in {data_dir}. Index not created.")
        return

    embeddings = get_embeddings()
    vs = FAISS.from_documents(docs, embeddings)
    os.makedirs(index_dir, exist_ok=True)
    vs.save_local(index_dir, "index_store")
    print(f"[Success] Indexed {len(docs)} documents. Saved index to {index_dir}")

def get_retriever(index_dir: str, k: int = 5):
    """Loads the FAISS index and returns a retriever."""
    if not os.path.exists(index_dir):
        print(f"[Error] Index directory not found: {index_dir}")
        return None
        
    embeddings = get_embeddings()
    vs = FAISS.load_local(index_dir, embeddings, allow_dangerous_deserialization=True, index_name="index_store")
    return vs.as_retriever(search_kwargs={"k": k})

def match(resume_path: str, jd_path: str, index_dir: str, k: int = 5) -> dict:
    """Performs the RAG-based matching between a resume and a job description using LCEL."""
    resume_text = extract_and_clean(resume_path)
    jd_text = extract_and_clean(jd_path)

    if not resume_text or not jd_text:
        return {"error": "Could not extract text from one or both files."}

    retriever = get_retriever(index_dir, k)
    if not retriever:
        return {"error": "Could not load the FAISS index."}

    # Setup LCEL RAG Chain
    llm = ChatOpenAI(temperature=0, model_name="gpt-3.5-turbo-0125")
    prompt = ChatPromptTemplate.from_template(MATCHING_PROMPT_TEMPLATE)
    
    # Define the inputs (resume, job) and the retrieval step (context)
    rag_chain = (
        RunnablePassthrough.assign(
            context=(lambda x: x["resume"] + "\n" + x["job"]) | retriever,
        ) 
        | prompt 
        | llm 
        | SimpleJsonOutputParser()
    )

    # Invoke the chain
    try:
        result = rag_chain.invoke({"resume": resume_text, "job": jd_text})
        print(json.dumps(result, indent=2))
        return result
    except Exception as e:
        return {"error": f"LLM Chain execution failed: {e}"}

# Streamlit Application
def run_ui(index_dir: str):
    """Initializes and runs the Streamlit web interface."""
    if not STREAMLIT_AVAILABLE:
        st.error("Streamlit is not installed.")
        return

    st.title("ResumeMatcher (RAG-Enabled)")
    r_file = st.file_uploader("Upload Resume", type=["pdf", "txt", "docx", "doc"])
    j_file = st.file_uploader("Upload Job Description", type=["pdf", "txt", "docx", "doc"])

    if r_file and j_file:
        # Use a dictionary to track temp files for easy cleanup
        temp_paths = {}
        try:
            for file, name in [(r_file, "r_path"), (j_file, "j_path")]:
                with tempfile.NamedTemporaryFile(delete=False, suffix=pathlib.Path(file.name).suffix) as tmp:
                    tmp.write(file.read())
                    temp_paths[name] = tmp.name

            if st.button("Get Match Score", use_container_width=True):
                with st.spinner("Analyzing..."):
                    result = match(temp_paths["r_path"], temp_paths["j_path"], index_dir)
                    
                    if "error" in result:
                        st.error(result["error"])
                    elif "score" in result:
                        st.success(f"### Match Score: {result['score']}/100")
                        st.write("### Key Explanations:")
                        if isinstance(result.get("explanation"), list):
                            st.markdown("\n".join(f"- {point}" for point in result["explanation"]))
                        else:
                            st.json(result)
                    else:
                        st.warning("The LLM returned a non-standard response.")
                        st.json(result)

        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
        finally:
            # Clean up all temporary files
            for path in temp_paths.values():
                if os.path.exists(path):
                    os.unlink(path)

# Main Execution
def main():
    """Main function to handle command-line arguments."""
    parser = argparse.ArgumentParser(description="A RAG-based tool for matching resumes to job descriptions.")
    
    # MODIFIED LINE: Remove 'required=True' and set 'default="serve"'
    parser.add_argument("--action", choices=["index", "match", "serve"], default="serve") 
    
    parser.add_argument("--data-dir", default="./data")
    parser.add_argument("--index-dir", default="./index")
    parser.add_argument("--resume")
    parser.add_argument("--jd")
    args = parser.parse_args()

    if args.action == "index":
        index_documents(args.data_dir, args.index_dir)
    elif args.action == "match":
        if not args.resume or not args.jd:
            print("[Error] --resume and --jd required for 'match'.")
            exit(1)
        match(args.resume, args.jd, args.index_dir)
    elif args.action == "serve":
        if not STREAMLIT_AVAILABLE:
            print("Install streamlit. Then run: streamlit run your_script_name.py -- --action serve")
            exit(1)
        run_ui(args.index_dir)

if __name__ == "__main__":
    main()
